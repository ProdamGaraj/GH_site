"""
Скрипт для конвертации Markdown в Word (.docx)
Использование: python md_to_word.py
"""

import re
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
except ImportError:
    print("Установка python-docx...")
    import subprocess
    subprocess.check_call(["pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE


def md_to_docx(md_path: str, docx_path: str):
    """Конвертирует Markdown файл в Word документ"""
    
    # Читаем markdown
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Создаем документ
    doc = Document()
    
    # Настраиваем стили
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)
    
    # Убираем лишние пустые строки (оставляем максимум одну подряд)
    content = re.sub(r'\n{3,}', '\n\n', content)
    
    # Разбиваем на строки
    lines = content.split('\n')
    i = 0
    prev_was_empty = False
    
    while i < len(lines):
        line = lines[i]
        
        # Пустая строка - пропускаем, не добавляем пустые параграфы
        if not line.strip():
            prev_was_empty = True
            i += 1
            continue
        
        # Заголовок 1
        if line.startswith('# '):
            p = doc.add_heading(line[2:], level=1)
            prev_was_empty = False
            i += 1
            continue
        
        # Заголовок 2
        if line.startswith('## '):
            p = doc.add_heading(line[3:], level=2)
            prev_was_empty = False
            i += 1
            continue
        
        # Заголовок 3
        if line.startswith('### '):
            p = doc.add_heading(line[4:], level=3)
            prev_was_empty = False
            i += 1
            continue
        
        # Горизонтальная линия - пропускаем
        if line.strip() == '---':
            i += 1
            continue
        
        # Чекбокс
        if line.strip().startswith('- [ ]'):
            text = line.strip()[6:]
            p = doc.add_paragraph()
            p.add_run('☐ ').bold = True
            p.add_run(process_inline_formatting(text))
            prev_was_empty = False
            i += 1
            continue
        
        # Маркированный список
        if line.strip().startswith('- ') and not line.strip().startswith('- ['):
            text = line.strip()[2:]
            p = doc.add_paragraph(process_inline_formatting(text), style='List Bullet')
            prev_was_empty = False
            i += 1
            continue
        
        # Жирный текст как подзаголовок (вопрос)
        if line.startswith('**') and line.rstrip().endswith('**'):
            text = line.strip()[2:-2]
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(11)
            prev_was_empty = False
            i += 1
            continue
        
        # Строка для ответа - добавляем сразу после вопроса без лишних отступов
        if line.startswith('Ответ:'):
            p = doc.add_paragraph(line)
            prev_was_empty = False
            i += 1
            continue
        
        # Строки с подчёркиваниями - пропускаем
        if line.strip().startswith('_____'):
            i += 1
            continue
        
        # Обычный текст
        p = doc.add_paragraph(process_inline_formatting(line))
        prev_was_empty = False
        i += 1
    
    # Сохраняем
    doc.save(docx_path)
    print(f"✅ Создан файл: {docx_path}")


def process_inline_formatting(text: str) -> str:
    """Убирает markdown форматирование из текста"""
    # Убираем жирный текст
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    # Убираем курсив
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    # Убираем код
    text = re.sub(r'`(.+?)`', r'\1', text)
    return text


if __name__ == '__main__':
    # Пути к файлам
    md_file = Path(__file__).parent / 'sprints' / 'sprint-0-preparation.md'
    docx_file = Path(__file__).parent / 'sprints' / 'sprint-0-preparation.docx'
    
    if md_file.exists():
        md_to_docx(str(md_file), str(docx_file))
    else:
        print(f"❌ Файл не найден: {md_file}")
